from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = "KrishiMitra_SIH_Briefing_Document.docx"

NAVY = "163B5C"
GREEN = "1E6B52"
GOLD = "B77820"
LIGHT_GREEN = "E8F4EE"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
MUTED = "5A6573"
WHITE = "FFFFFF"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_borders(table, color="D8DEE5", size="6"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement('w:tblHeader')
    node.set(qn('w:val'), 'true')
    trPr.append(node)

def set_fixed_table_layout(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])

def set_font(run, size=11, color="000000", bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def para(doc, text="", size=11, color="000000", bold=False, italic=False, align=None, before=0, after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size, color, bold, italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), 10.5)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_font(p.add_run(text), 10.5)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_font(r, 16, NAVY, True)
    else:
        set_font(r, 12.5, GREEN if level == 2 else NAVY, True)
    return p

def add_callout(doc, title, body, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(table, [6.5])
    set_table_borders(table, fill, '8')
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), 11, GREEN, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_font(p2.add_run(body), 10.5, "24323F")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), font_size, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_font(p.add_run(str(text)), font_size, "24323F")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
styles['Normal']._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
styles['Normal'].font.size = Pt(11)

# Header/footer
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_p.paragraph_format.space_after = Pt(0)
set_font(header_p.add_run('KRISHIMITRA | SIH BRIEFING DOCUMENT'), 8.5, MUTED, True)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(0)
set_font(footer_p.add_run('ZeroTheory | KrishiMitra | Prepared for Smart India Hackathon 2026'), 8, MUTED)

# Cover
para(doc, 'SMART INDIA HACKATHON 2026', 10, GOLD, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
para(doc, 'KrishiMitra', 30, NAVY, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(doc, 'AI-powered farm decision support for Indian farmers', 15, GREEN, False, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
add_callout(doc, 'One-line pitch', 'KrishiMitra is a multilingual AI farm companion that helps Indian farmers detect crop and soil problems early and take the right next step through simple photo and voice-based guidance.', LIGHT_GREEN)
para(doc, 'SIH Pitch, MVP, Market Research & Business Model', 14, NAVY, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=6)
para(doc, 'Prepared by Team Alvengers | ZeroTheory | 30 August 2026', 10.5, MUTED, False, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
para(doc, 'Document purpose: a judge-ready narrative for SIH presentations, Q&A rounds, and partner discussions. Revenue figures are planning estimates, not audited forecasts.', 9.5, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
doc.add_page_break()

add_heading(doc, '1. Executive Summary')
para(doc, 'KrishiMitra is a mobile-first agricultural assistant designed for Indian farmers. It brings crop and soil image analysis, multilingual chat, voice-friendly interaction, mandi information, fertilizer guidance, government schemes, and farmer-support resources into one focused application. The product’s role is decision support: help the farmer spot issues earlier, understand the next action, and connect to relevant services. It does not replace qualified agronomists or agricultural officers.')
add_callout(doc, 'Why now', 'India’s agriculture and allied activities contribute nearly one-fifth of national income and account for 46.1% of the workforce. The scale of the sector, combined with fragmented farm guidance, creates a strong need for accessible, locally relevant decision-support tools. [S1]', LIGHT_BLUE)

add_heading(doc, '2. Problem Statement')
add_bullet(doc, 'Crop diseases, pests, nutrient deficiencies, and soil issues can become more costly when identified late.')
add_bullet(doc, 'Farmers often need to consult multiple, fragmented sources for crop advice, weather, mandi information, schemes, and support.')
add_bullet(doc, 'General AI tools are not built around farm workflows, local languages, or the practical next actions a farmer needs.')
add_bullet(doc, 'Low digital familiarity and uneven connectivity can make technical, text-heavy tools difficult to use.')

add_heading(doc, '3. Unique Selling Proposition (USP)')
add_callout(doc, 'USP statement', 'KrishiMitra turns complex AI into a simple, local-language farm companion: a farmer can upload a crop or soil photo, speak or type a question, and receive clear next-step guidance connected to farm information and support.', LIGHT_GREEN)
add_table(doc, ['USP pillar', 'What it means for the farmer'], [
    ('Focused rather than generic', 'The workflow is built around crop health, soil, fertilizer, market information, schemes, and support rather than broad general-purpose prompting.'),
    ('Photo + voice + local language', 'Farmers can submit a crop or soil image and use voice or text in a familiar Indian language.'),
    ('Action, not only diagnosis', 'Results include diagnosis, confidence, recommendations, treatment steps, and preventive measures.'),
    ('One connected journey', 'The same platform links analysis with fertilizer guidance, mandi information, schemes, history, and farmer-support resources.'),
    ('Safety-oriented decision support', 'Confidence levels and expert-verification advice keep the product positioned as support, not a substitute for qualified agronomy advice.')
], [1.8, 4.7])

add_heading(doc, '4. MVP: What We Can Demonstrate Today')
add_table(doc, ['MVP capability', 'Demonstration outcome'], [
    ('Crop image analysis', 'Upload a crop photo and receive crop type, possible diagnosis, status, confidence, recommendations, treatment steps, and preventive measures.'),
    ('Soil image guidance', 'Upload a soil image and receive soil type, condition, estimated pH description, fertility assessment, and improvement suggestions.'),
    ('Multilingual experience', 'The interface supports 14 Indian languages, with selected-language AI analysis and voice-input locale support.'),
    ('Contextual AI chat', 'Ask farming questions by text or voice and receive guidance tailored to the farmer’s saved region and crop context.'),
    ('Farm-planning tools', 'Use fertilizer guidance, crop calendar, regional weather, mandi-price information, scheme discovery, and farmer-support resources.'),
    ('History & continuity', 'Save completed analysis results and chat activity for later reference.')
], [1.8, 4.7])
para(doc, 'MVP boundary: live AI image analysis requires connectivity and image quality affects confidence. Critical pesticide, fertilizer, or disease decisions should be verified with a qualified local expert, KVK, or agricultural officer.', 9.5, MUTED, italic=True)

add_heading(doc, '5. Market Research: India Agriculture & Agritech Opportunity')
para(doc, 'The market case is not based only on a large farmer population. KrishiMitra targets a measurable operating need: easier access to crop-health guidance, local-language interaction, and farm services. The most credible early buyer is often an institution serving many farmers, such as an FPO, NGO, cooperative, or public programme.')
add_table(doc, ['Research finding', 'Why it matters to KrishiMitra'], [
    ('Agriculture and allied activities contribute nearly one-fifth of India’s national income and employ 46.1% of the workforce. [S1]', 'A farmer-support product addresses a large, economically central user group; the value proposition must focus on practical, low-friction use.'),
    ('Agriculture and allied GVA at current prices rose to about ₹48.78 lakh crore in 2023-24. [S2]', 'The sector has national economic significance; solutions that reduce avoidable loss or improve decisions have a meaningful adoption rationale.'),
    ('The Government of India reported 10,000 registered FPOs in 2025, connecting about 30 lakh farmers; around 40% were women. [S3]', 'FPOs create a concrete institutional distribution and B2B buyer channel for pilots, training, dashboards, and group support.'),
    ('IBEF reports India’s smart-agriculture market at US$860.7 million in 2025, with an 18.93% projected CAGR to 2034. [S4]', 'This is a directional third-party market estimate, not KrishiMitra revenue. It supports the case that precision and digital agriculture are growing categories.'),
    ('Government policy is actively building FPO and digital-market infrastructure, including e-NAM, ONDC, and GeM onboarding support. [S5]', 'Partnerships and integrations can be more viable than building every data layer independently.')
], [2.55, 3.95], 9)
add_callout(doc, 'Research interpretation', 'The addressable opportunity should be validated from pilots, not inferred from the entire agriculture or smart-agriculture market. Our initial beachhead is FPOs, NGOs, and local farmer groups; this gives KrishiMitra an acquisition channel, a training partner, and a path to paid institutional deployment.', LIGHT_BLUE)

add_heading(doc, '6. Business Model')
para(doc, 'KrishiMitra uses a farmer-first freemium model. Essential guidance remains free, while advanced convenience, institutional deployment, and verified service connections create sustainable revenue.')
add_table(doc, ['Revenue stream', 'Offer', 'Indicative pricing', 'Strategic role'], [
    ('Free farmer tier', 'Basic guidance, limited analysis, multilingual access, schemes and support information.', '₹0', 'Build trust, reach, and repeat use.'),
    ('Premium farmer tier', 'More analyses, detailed reports, alerts, saved history, and tailored planning.', '₹79-₹99/month', 'Recurring B2C revenue; optional, not a barrier to critical guidance.'),
    ('One-time detailed report', 'A detailed crop or soil analysis for farmers who do not want a subscription.', '₹20/report', 'Low-friction pay-per-use option.'),
    ('FPO / NGO dashboard', 'Multi-farmer support, onboarding, programme reporting, and aggregate insights subject to consent and privacy controls.', '₹5,000-₹7,500/month', 'Primary early recurring revenue channel.'),
    ('Government / CSR pilot', 'Configured deployment, training, support, and evaluation for a district, programme, or farmer cohort.', '₹2.5-₹6 lakh/year', 'High-trust, partnership-led growth channel.'),
    ('Verified referrals - future', 'Transparent referrals to soil labs, insurance, input suppliers, equipment rental, or experts.', 'Commission-based', 'Only after verification and explicit disclosure; never compromise advisory trust.')
], [1.35, 2.55, 1.1, 1.5], 8.8)

add_heading(doc, '7. Estimated Revenue & Path to Profitability')
para(doc, 'The following are planning scenarios for SIH discussion. They are not audited forecasts and must be validated through pilots, actual acquisition costs, AI inference costs, support costs, and retention data.')
add_table(doc, ['Scenario', 'Assumptions', 'Estimated monthly revenue', 'Estimated annual revenue'], [
    ('Pilot', '60 premium farmers at ₹79; 300 reports at ₹20; 5 FPO/NGO partners at ₹5,000; one ₹2.5 lakh annual pilot contract.', '₹56,573', '₹6.79 lakh'),
    ('Expansion', '1,250 premium farmers at ₹99; 1,000 reports at ₹20; 20 institutional partners at ₹7,500; ₹6 lakh/year in programme contracts.', '₹3,43,750', '₹41.25 lakh')
], [1.0, 3.6, 1.0, 0.9], 9)
para(doc, 'Pilot calculation: (60 × ₹79) + (300 × ₹20) + (5 × ₹5,000) + (₹2,50,000 ÷ 12) = ₹56,573/month average. Expansion calculation: (1,250 × ₹99) + (1,000 × ₹20) + (20 × ₹7,500) + (₹6,00,000 ÷ 12) = ₹3,43,750/month average.', 9, MUTED, italic=True)
add_heading(doc, 'How profitability will be achieved', 2)
add_numbered(doc, 'Use FPO, NGO, and programme contracts to cover core platform, onboarding, and support costs before relying on individual farmer subscriptions.')
add_numbered(doc, 'Keep free features lightweight; apply AI image-analysis usage limits and offer premium or pay-per-report access for higher-cost usage.')
add_numbered(doc, 'Measure API inference cost per scan, active-user retention, support cost, and conversion rates in every pilot.')
add_numbered(doc, 'Expand only after the contribution margin from subscriptions and institutional contracts exceeds infrastructure, AI, field-training, and customer-support costs.')
add_callout(doc, 'Profitability principle', 'KrishiMitra becomes sustainable when recurring institutional revenue and voluntary premium revenue exceed the real cost of AI inference, cloud hosting, field onboarding, support, and compliance. The pilot should be used to establish those actual unit economics.', LIGHT_GREEN)

add_heading(doc, '8. Go-to-Market Plan')
add_table(doc, ['Phase', 'Action', 'Success measure'], [
    ('Pilot - 0 to 6 months', 'Run village demos with FPOs, NGOs, agriculture colleges, KVKs, panchayats, and farmer groups. Train farmer champions.', 'First 1,000 registered farmers; repeat analysis and voice use; farmer feedback.'),
    ('Validate - 6 to 12 months', 'Offer institutional dashboards and selected premium features. Improve recommendations through expert review and pilot data.', 'Paid institutional pilots; retention; documented accuracy and satisfaction measures.'),
    ('Scale - 12+ months', 'Expand state-wise through FPO networks, scheme-awareness campaigns, local-language content, and verified partners.', 'Recurring contracts, lower acquisition cost, and a positive contribution margin.')
], [1.35, 3.7, 1.45], 9)

add_heading(doc, '9. Judge Q&A Cheat Sheet')
qa = [
    ('What problem are you solving?', 'KrishiMitra helps farmers identify crop and soil issues earlier and access practical, local-language farm guidance without searching across multiple sources.'),
    ('What makes you different?', 'We are focused on the farmer workflow: photo or voice input, local language, clear next steps, and links to fertilizer guidance, mandi information, schemes, and support.'),
    ('What works in the MVP?', 'Crop and soil image guidance, multilingual chat and voice input, fertilizer guidance, mandi information, schemes, support resources, and saved history are demonstrable.'),
    ('How do you keep AI advice safe?', 'We show confidence levels, request clearer images for uncertain cases, and position the app as decision support. Serious cases should be verified with qualified experts.'),
    ('How do you handle unclear images?', 'A low-confidence result prompts the farmer to upload a clearer close-up image and seek expert verification rather than treating the output as final.'),
    ('How will you validate accuracy?', 'We will test labelled images across crops, regions, lighting, and disease stages, then review results with agriculture experts, KVKs, and farmer groups.'),
    ('How do low-digital-literacy users use it?', 'The app is mobile-first, supports 14 Indian languages, and offers voice interaction so farmers do not need to type technical questions.'),
    ('How do you address poor connectivity?', 'We keep the experience lightweight and show offline status clearly. Live AI image analysis needs connectivity, so the app must state that transparently.'),
    ('How will you get your first 1,000 farmers?', 'Through local workshops, FPOs, NGOs, KVKs, panchayat farmer groups, trained farmer champions, WhatsApp groups, and regional outreach.'),
    ('Why would FPOs or NGOs pay?', 'They can support more farmers with a ready-to-use, local-language decision-support layer instead of building AI, image analysis, and farmer workflows from scratch.'),
    ('How do you scale across states?', 'The same core platform adapts through the farmer’s language, state, crop context, and modular local-information services; we scale via regional partner networks.'),
    ('How do you earn revenue?', 'Essential guidance stays free. Revenue comes from optional premium tools, one-time reports, and institutional subscriptions and deployments.'),
    ('What is your biggest limitation?', 'Image quality and real-world validation. Our next phase is expert-reviewed datasets, official data integrations, and pilots across crops and regions.'),
    ('What will you do after SIH?', 'Run district pilots, measure outcomes, validate recommendations with experts, strengthen verified data integrations, and expand through institutional partners.')
]
for question, answer in qa:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run('Q. ' + question), 10.5, NAVY, True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.15)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.line_spacing = 1.15
    set_font(p2.add_run('A. ' + answer), 10, '24323F')

add_heading(doc, '10. Research Sources')
sources = [
    ('[S1]', 'Government of India, Economic Survey 2025-26, Chapter 6: Agriculture and Food Management. Agriculture and allied activities contribute nearly one-fifth of national income and account for 46.1% of the workforce. https://www.indiabudget.gov.in/economicsurvey/doc/echapter.pdf'),
    ('[S2]', 'Ministry of Statistics and Programme Implementation (MoSPI), Statistical Report on Value of Output from Agriculture and Allied Sectors, 2011-12 to 2023-24 (2025). Reports agriculture and allied GVA at current prices of about ₹48.78 lakh crore in 2023-24. https://www.mospi.gov.in/sites/default/files/press_release/Press%20release%20for%20Brochure%202025_0.pdf'),
    ('[S3]', 'Press Information Bureau, Ministry of Agriculture & Farmers Welfare, “10,000 FPOs Achieved under Government’s Flagship Scheme,” 28 February 2025. https://www.pib.gov.in/PressReleasePage.aspx?PRID=2106913&lang=1&reg=3'),
    ('[S4]', 'India Brand Equity Foundation (IBEF), “Indian Agriculture Sector, Farming in India,” accessed 30 August 2026. Market-size forecast cited as a third-party estimate and should be independently verified before investment use. https://www.ibef.org/industry/agriculture-india'),
    ('[S5]', 'Press Information Bureau, Ministry of Agriculture & Farmers Welfare, “Formation and Promotion of FPOs,” 8 August 2025. https://www.pib.gov.in/PressReleasePage.aspx?PRID=2154174&lang=1&reg=3')
]
for label, source in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(label + ' '), 9, NAVY, True)
    set_font(p.add_run(source), 9, "3C4652")

para(doc, 'Note on research use: Government sources support the policy and sector facts in this document. IBEF market-size projections are directional market research, not a guarantee of KrishiMitra demand, revenue, or profitability.', 8.7, MUTED, italic=True, before=8, after=0)

doc.core_properties.title = 'KrishiMitra SIH Briefing Document'
doc.core_properties.subject = 'SIH pitch, MVP, market research, revenue model and Q&A'
doc.core_properties.author = 'Team Alvengers / ZeroTheory'
doc.save(OUT)
print(OUT)
